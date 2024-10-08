import os
import whisper
import threading
from docx import Document
from tkinter import Tk, Label, Button, filedialog, StringVar, ttk, BooleanVar, Checkbutton

# Variável global para o cancelamento da transcrição
cancelar_transcricao = False

# Extensões de áudio suportadas
EXTENSOES_SUPORTADAS = ('.mp3', '.wav', '.flac', '.m4a', '.ogg', '.mpeg')

# Função para carregar o modelo Whisper
def carregar_modelo(modelo_escolhido):
    print(f"Carregando o modelo '{modelo_escolhido}'...")
    return whisper.load_model(modelo_escolhido)

# Função para transcrever o áudio
def transcrever_audio(modelo, caminho_audio, progresso_var, status_label, index, total):
    global cancelar_transcricao
    try:
        if cancelar_transcricao:
            status_label.config(text="Transcrição cancelada.")
            return False
        status_label.config(text=f"Transcrevendo arquivo {index + 1} de {total}: {os.path.basename(caminho_audio)}")
        resposta = modelo.transcribe(caminho_audio)
        if cancelar_transcricao:
            status_label.config(text="Transcrição cancelada.")
            return False
        texto_transcrito = resposta['text']
        salvar_transcricao(texto_transcrito, caminho_audio)
        progresso_var.set(100)
        return True
    except RuntimeError as e:
        if "reshape tensor" in str(e):
            status_label.config(text="Erro: arquivo corrompido ou incompatível")
        else:
            status_label.config(text=f"Erro ao transcrever: {e}")
        return False
    except Exception as e:
        status_label.config(text=f"Erro inesperado: {e}")
        return False

# Função para salvar a transcrição
def salvar_transcricao(texto_transcrito, caminho_audio):
    nome_arquivo, _ = os.path.splitext(os.path.basename(caminho_audio))
    diretorio_arquivo = os.path.dirname(caminho_audio)
    caminho_arquivo_docx = os.path.join(diretorio_arquivo, f"{nome_arquivo}_transcrito.docx")
    doc = Document()
    doc.add_heading(f"Transcrição de {nome_arquivo}", level=1)
    doc.add_paragraph(texto_transcrito)
    doc.save(caminho_arquivo_docx)
    print(f"Transcrição salva no arquivo: {caminho_arquivo_docx}")

# Função para iniciar transcrição de um único arquivo
def transcrever_individualmente(modelo_escolhido, progresso_var, root, status_label, botao_cancelar):
    def run_transcricao():
        global cancelar_transcricao
        cancelar_transcricao = False  # Reseta o cancelamento no início
        root.withdraw()
        caminho_audio = filedialog.askopenfilename(title="Selecione o arquivo de áudio", filetypes=[("Arquivos de áudio", "*.mp3 *.wav *.flac *.m4a *.ogg *.mpeg")])
        root.deiconify()
        if not caminho_audio:
            status_label.config(text="Nenhum arquivo selecionado.")
            return
        botao_cancelar.config(state="normal")  # Habilita o botão cancelar após a seleção
        modelo = carregar_modelo(modelo_escolhido)
        status_label.config(text="Transcrevendo 1 arquivo")
        sucesso = transcrever_audio(modelo, caminho_audio, progresso_var, status_label, 0, 1)  # Executa sequencialmente
        if sucesso and not cancelar_transcricao:
            status_label.config(text="A transcrição foi realizada com sucesso.")
        botao_cancelar.config(state="disabled")  # Desabilita o botão cancelar após a transcrição ou cancelamento
    
    threading.Thread(target=run_transcricao).start()

# Função para iniciar transcrição de múltiplos arquivos (lote), com opção de subpastas
def processar_em_lote(modelo_escolhido, progresso_var, root, status_label, botao_cancelar, incluir_subpastas):
    def run_lote():
        global cancelar_transcricao
        cancelar_transcricao = False  # Reseta o cancelamento no início
        root.withdraw()
        pasta = filedialog.askdirectory(title="Selecione a pasta com os arquivos de áudio")
        root.deiconify()
        if not pasta:
            status_label.config(text="Nenhuma pasta selecionada.")
            return
        botao_cancelar.config(state="normal")  # Habilita o botão cancelar após a seleção
        modelo = carregar_modelo(modelo_escolhido)

        # Coleta arquivos de áudio na pasta e subpastas se habilitado
        arquivos_audio = []
        if incluir_subpastas.get():
            for root_dir, _, files in os.walk(pasta):
                arquivos_audio.extend([os.path.join(root_dir, f) for f in files if f.endswith(EXTENSOES_SUPORTADAS)])
        else:
            arquivos_audio = [os.path.join(pasta, f) for f in os.listdir(pasta) if f.endswith(EXTENSOES_SUPORTADAS)]
        
        if not arquivos_audio:
            status_label.config(text="Nenhum arquivo de áudio encontrado.")
            botao_cancelar.config(state="disabled")  # Desabilita o botão cancelar
            return

        progresso_total = 0
        progresso_por_arquivo = 100 / len(arquivos_audio)
        
        # Processa cada arquivo sequencialmente, mostrando o progresso
        for index, arquivo in enumerate(arquivos_audio):
            if cancelar_transcricao:
                status_label.config(text="Transcrição cancelada.")
                break
            transcrever_audio(modelo, arquivo, progresso_var, status_label, index, len(arquivos_audio))
            progresso_total += progresso_por_arquivo
            progresso_var.set(progresso_total)

        if not cancelar_transcricao:
            status_label.config(text=f"As {len(arquivos_audio)} transcrições foram realizadas com sucesso.")
        botao_cancelar.config(state="disabled")  # Desabilita o botão cancelar após a transcrição ou cancelamento
    
    threading.Thread(target=run_lote).start()

# Função para cancelar a transcrição
def cancelar_processo():
    global cancelar_transcricao
    cancelar_transcricao = True
    print("Processo de transcrição cancelado.")

# Função principal da interface gráfica
def criar_interface():
    root = Tk()
    root.title("Transcrição de Áudio com Whisper")
    root.geometry("450x400")
    root.resizable(True, True)

    progresso_var = StringVar(value=0)
    Label(root, text="Escolha o modelo de transcrição Whisper:", padx=20, pady=10).grid(row=0, column=0, columnspan=2)
    modelo_escolhido = StringVar(value="tiny")
    ttk.Combobox(root, textvariable=modelo_escolhido, values=["tiny", "base", "small", "medium", "large", "turbo"], width=15).grid(row=1, column=0, columnspan=2, padx=20, pady=5)
    
    # Status Label
    status_label = Label(root, text="", padx=20, pady=5, fg="blue", font=("Arial", 12), wraplength=500)
    status_label.grid(row=3, column=0, columnspan=2, padx=20, pady=5)

    # Opção de transcrição individual e em lote
    Button(root, text="Transcrição Individual", command=lambda: iniciar_transcricao(modelo_escolhido.get(), progresso_var, root, status_label, botao_cancelar), width=25).grid(row=2, column=0, padx=20, pady=10)
    Button(root, text="Transcrição em Lote", command=lambda: iniciar_transcricao_em_lote(modelo_escolhido.get(), progresso_var, root, status_label, botao_cancelar, incluir_subpastas), width=25).grid(row=2, column=1, padx=20, pady=10)
    
    # Checkbox para incluir subpastas
    incluir_subpastas = BooleanVar()
    Checkbutton(root, text="Incluir Subpastas", variable=incluir_subpastas).grid(row=4, column=0, columnspan=2, pady=5)

    # Barra de progresso
    progresso_barra = ttk.Progressbar(root, length=300, mode='determinate', variable=progresso_var)
    progresso_barra.grid(row=5, column=0, columnspan=2, padx=20, pady=20)

    # Botão de cancelar e fechar
    botao_cancelar = Button(root, text="Cancelar Transcrição", command=cancelar_processo, width=25, state="disabled")
    botao_cancelar.grid(row=6, column=0, padx=20, pady=10)
    botao_fechar = Button(root, text="Fechar", command=root.quit, width=25)
    botao_fechar.grid(row=6, column=1, padx=20, pady=10)

    root.mainloop()

def iniciar_transcricao(modelo_escolhido, progresso_var, root, status_label, botao_cancelar):
    transcrever_individualmente(modelo_escolhido, progresso_var, root, status_label, botao_cancelar)

def iniciar_transcricao_em_lote(modelo_escolhido, progresso_var, root, status_label, botao_cancelar, incluir_subpastas):
    processar_em_lote(modelo_escolhido, progresso_var, root, status_label, botao_cancelar, incluir_subpastas)

if __name__ == "__main__":
    criar_interface()
