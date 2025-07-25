import os
import whisper
import threading
import logging
import time
import subprocess
import platform
import tempfile
import json
from enum import Enum
from datetime import datetime
from docx import Document
from tkinter import Tk, Label, Button, filedialog, StringVar, ttk, BooleanVar, Checkbutton, Text, Scrollbar, Frame, \
    NORMAL, DISABLED, messagebox, IntVar

# --- Importação condicional do AudioSegment e tratativa de erro ---
try:
    from pydub import AudioSegment

    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False
    logging.error("A biblioteca 'pydub' não foi encontrada. Por favor, instale-a: pip install pydub")
    messagebox.showerror("Erro de Dependência",
                         "A biblioteca 'pydub' não foi encontrada. Por favor, instale-a usando 'pip install pydub'.")
except Exception as e:
    PYDUB_AVAILABLE = False
    logging.error(
        f"Erro ao importar 'pydub' ou 'AudioSegment'. Verifique se 'ffmpeg' está instalado e no PATH do sistema. Erro: {e}")
    messagebox.showerror("Erro de Dependência",
                         f"Erro ao carregar 'pydub'. Certifique-se de que 'ffmpeg' está instalado e configurado corretamente no seu sistema (PATH). Erro: {e}")

# --- Configuração do Logger ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('transcricao.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)


# --- Enums para melhor organização ---
class AudioExtension(Enum):
    MP3 = '.mp3'
    WAV = '.wav'
    FLAC = '.flac'
    M4A = '.m4a'
    OGG = '.ogg'
    MPEG = '.mpeg'
    WEBM = '.webm'
    AAC = '.aac'


class WhisperModel(Enum):
    TINY = "tiny"
    BASE = "base"
    SMALL = "small"
    MEDIUM = "medium"
    LARGE = "large"
    TURBO = "turbo"


class TranscricaoAudio:
    MODELOS_DESCRICAO = {
        WhisperModel.TINY.value: "Tiny: O modelo mais leve e rápido, ideal para tarefas rápidas com precisão básica; requer poucos recursos.",
        WhisperModel.BASE.value: "Base: Equilíbrio entre velocidade e precisão moderada, funciona bem em dispositivos com recursos limitados.",
        WhisperModel.SMALL.value: "Small: Boa precisão e velocidade razoável, recomendado para uso geral; requer recursos moderados.",
        WhisperModel.MEDIUM.value: "Medium: Muito preciso, excelente para diferentes sotaques e contextos, ideal para máquinas com boa capacidade.",
        WhisperModel.LARGE.value: "Large: Máxima precisão, melhor qualidade de transcrição; exige muitos recursos, ideal para máquinas potentes.",
        WhisperModel.TURBO.value: "Turbo: Otimizado para máxima velocidade com alta precisão, indicado para servidores ou estações de alta potência."
    }

    IDIOMAS_WHISPER = {
        "auto": "Detectar automaticamente",
        "pt": "Português",
        "en": "Inglês",
        "es": "Espanhol",
        "fr": "Francês",
        "de": "Alemão",
        "it": "Italiano",
        "ja": "Japonês",
        "ko": "Coreano",
        "zh": "Chinês",
        "ru": "Russo",
        "ar": "Árabe"
    }

    CONFIG_FILE = "config_transcricao.json"

    def __init__(self):
        self.cancel_event = threading.Event()
        self.pause_event = threading.Event()
        self.total_bytes = 0
        self.processed_bytes = 0
        self.start_time = 0
        self.modelo_carregado = None
        self.modelo_carregado_nome = None
        self.estatisticas = {
            'arquivos_processados': 0,
            'tempo_total_processamento': 0,
            'erros': 0,
            'sucessos': 0
        }

        self.root = Tk()
        self._inicializar_variaveis()
        self._carregar_configuracoes()
        self._configurar_interface()

        # Verificar se pydub está disponível, se não, desabilitar botões de início
        if not PYDUB_AVAILABLE:
            messagebox.showerror("Erro de Configuração",
                                 "FFmpeg ou pydub não estão configurados corretamente. O aplicativo pode não funcionar. Por favor, consulte o log para mais detalhes.")
            self._set_transcription_controls_state(False)

    def _inicializar_variaveis(self):
        self.modelo_escolhido = StringVar(value=WhisperModel.TURBO.value)
        self.progresso_var = StringVar(value="0")
        self.incluir_subpastas = BooleanVar()
        self.formato_saida = StringVar(value="docx")
        self.idioma_escolhido = StringVar(value="auto")
        self.temperatura = StringVar(value="0.0")
        self.segmento_duracao = IntVar(value=30)
        self.incluir_timestamps = BooleanVar()
        self.pasta_saida_personalizada = StringVar()

    def _carregar_configuracoes(self):
        """Carrega configurações salvas do arquivo JSON"""
        try:
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)

                self.modelo_escolhido.set(config.get('modelo', WhisperModel.TURBO.value))
                self.formato_saida.set(config.get('formato', 'docx'))
                self.idioma_escolhido.set(config.get('idioma', 'auto'))
                self.temperatura.set(config.get('temperatura', '0.0'))
                self.segmento_duracao.set(config.get('segmento_duracao', 30))
                self.incluir_subpastas.set(config.get('incluir_subpastas', False))
                self.incluir_timestamps.set(config.get('incluir_timestamps', False))
                self.pasta_saida_personalizada.set(config.get('pasta_saida', ''))

                logging.info("Configurações carregadas com sucesso")
        except Exception as e:
            logging.error(f"Erro ao carregar configurações: {e}")

    def _salvar_configuracoes(self):
        """Salva as configurações atuais no arquivo JSON"""
        try:
            config = {
                'modelo': self.modelo_escolhido.get(),
                'formato': self.formato_saida.get(),
                'idioma': self.idioma_escolhido.get(),
                'temperatura': self.temperatura.get(),
                'segmento_duracao': self.segmento_duracao.get(),
                'incluir_subpastas': self.incluir_subpastas.get(),
                'incluir_timestamps': self.incluir_timestamps.get(),
                'pasta_saida': self.pasta_saida_personalizada.get()
            }

            with open(self.CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2, ensure_ascii=False)

            logging.info("Configurações salvas com sucesso")
        except Exception as e:
            logging.error(f"Erro ao salvar configurações: {e}")

    def _configurar_interface(self):
        self.root.title("Transcrição de Áudio com Whisper - Versão Avançada")
        self.root.geometry("800x700")
        self.root.resizable(True, True)

        # Criar notebook para abas
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=10)

        # Aba Principal
        self.frame_principal = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_principal, text="Transcrição")
        self._configurar_aba_principal()

        # Aba Configurações
        self.frame_config = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_config, text="Configurações")
        self._configurar_aba_configuracoes()

        # Aba Estatísticas
        self.frame_stats = ttk.Frame(self.notebook)
        self.notebook.add(self.frame_stats, text="Estatísticas")
        self._configurar_aba_estatisticas()

    def _configurar_aba_principal(self):
        frame = self.frame_principal

        # Configurar grid
        frame.grid_rowconfigure(7, weight=1)
        frame.grid_rowconfigure(18, weight=1)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)

        # Caixa de Seleção do Modelo
        ttk.Label(frame, text="Modelo Whisper:", font=("Arial", 10, "bold")).grid(row=0, column=0, columnspan=2,
                                                                                  pady=(10, 5))
        self.combobox_modelo = ttk.Combobox(frame, textvariable=self.modelo_escolhido,
                                            values=[model.value for model in WhisperModel], width=15, state="readonly")
        self.combobox_modelo.grid(row=1, column=0, columnspan=2, padx=20, pady=5)
        self.combobox_modelo.bind("<<ComboboxSelected>>", self._atualizar_descricao_modelo)

        # Explicação sobre o Modelo
        self.modelo_explicacao = ttk.Label(frame, text=self.MODELOS_DESCRICAO[self.modelo_escolhido.get()],
                                           wraplength=700, justify="left")
        self.modelo_explicacao.grid(row=2, column=0, columnspan=2, padx=20, pady=5)

        # Status do modelo carregado
        self.status_modelo = ttk.Label(frame, text="Nenhum modelo carregado", foreground="red")
        self.status_modelo.grid(row=3, column=0, columnspan=2, pady=5)

        # Frame para botões principais
        frame_botoes = ttk.Frame(frame)
        frame_botoes.grid(row=4, column=0, columnspan=2, pady=10)

        self.btn_transcricao_individual = ttk.Button(frame_botoes, text="Transcrição Individual",
                                                     command=self.iniciar_transcricao, width=25)
        self.btn_transcricao_individual.pack(side="left", padx=5)

        self.btn_transcricao_em_lote = ttk.Button(frame_botoes, text="Transcrição em Lote",
                                                  command=self.iniciar_transcricao_em_lote, width=25)
        self.btn_transcricao_em_lote.pack(side="left", padx=5)

        # Opções rápidas
        frame_opcoes = ttk.LabelFrame(frame, text="Opções Rápidas", padding=10)
        frame_opcoes.grid(row=5, column=0, columnspan=2, padx=20, pady=10, sticky="ew")

        ttk.Checkbutton(frame_opcoes, text="Incluir Subpastas", variable=self.incluir_subpastas).pack(side="left",
                                                                                                      padx=10)
        ttk.Checkbutton(frame_opcoes, text="Incluir Timestamps", variable=self.incluir_timestamps).pack(side="left",
                                                                                                        padx=10)

        # Formato de saída
        ttk.Label(frame_opcoes, text="Formato:").pack(side="left", padx=(20, 5))
        ttk.Combobox(frame_opcoes, textvariable=self.formato_saida, values=["docx", "txt", "markdown", "srt"], width=10,
                     state="readonly").pack(side="left")

        # Barra de Progresso
        self.progresso_barra = ttk.Progressbar(frame, length=700, mode='determinate')
        self.progresso_barra.grid(row=6, column=0, columnspan=2, padx=20, pady=20, sticky="ew")

        # Labels de progresso
        self.progresso_text_label = ttk.Label(frame, text="", wraplength=700, justify="left")
        self.progresso_text_label.grid(row=7, column=0, columnspan=2, padx=20, pady=5, sticky="nw")

        self.eta_label = ttk.Label(frame, text="", foreground="blue")
        self.eta_label.grid(row=8, column=0, columnspan=2, padx=20, pady=2)

        # Frame para botões de controle
        frame_controle = ttk.Frame(frame)
        frame_controle.grid(row=9, column=0, columnspan=2, pady=10)

        self.botao_cancelar = ttk.Button(frame_controle, text="Cancelar", command=self.cancelar_processo, width=15,
                                         state=DISABLED)
        self.botao_cancelar.pack(side="left", padx=5)

        self.botao_pausar = ttk.Button(frame_controle, text="Pausar", command=self.pausar_processo, width=15,
                                       state=DISABLED)
        self.botao_pausar.pack(side="left", padx=5)

        # Frame para botões secundários
        frame_secundario = ttk.Frame(frame)
        frame_secundario.grid(row=10, column=0, columnspan=2, pady=10)

        self.botao_detalhes = ttk.Button(frame_secundario, text="Ver Detalhes", command=self._alternar_painel_detalhes,
                                         width=15)
        self.botao_detalhes.pack(side="left", padx=5)

        ttk.Button(frame_secundario, text="Limpar Log", command=self._limpar_detalhes, width=15).pack(side="left",
                                                                                                      padx=5)

        self.botao_fechar = ttk.Button(frame_secundario, text="Fechar", command=self._fechar_aplicacao, width=15)
        self.botao_fechar.pack(side="left", padx=5)

        # Painel de Detalhes
        self.painel_detalhes = ttk.LabelFrame(frame, text="Detalhes da Transcrição", padding=5)
        self.detalhes_text = Text(self.painel_detalhes, height=8, width=80, state=DISABLED, wrap="word")
        scrollbar_detalhes = ttk.Scrollbar(self.painel_detalhes, command=self.detalhes_text.yview)
        self.detalhes_text['yscrollcommand'] = scrollbar_detalhes.set

        self.detalhes_text.pack(side="left", fill="both", expand=True)
        scrollbar_detalhes.pack(side="right", fill="y")

        self.painel_detalhes.grid(row=18, column=0, columnspan=2, padx=10, pady=5, sticky="nsew")
        self.painel_detalhes.grid_remove()

    def _configurar_aba_configuracoes(self):
        frame = self.frame_config

        # Configurações de Transcrição
        config_frame = ttk.LabelFrame(frame, text="Parâmetros de Transcrição", padding=10)
        config_frame.pack(fill="x", padx=10, pady=10)

        # Idioma
        row = 0
        ttk.Label(config_frame, text="Idioma:").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        idioma_combo = ttk.Combobox(config_frame, textvariable=self.idioma_escolhido,
                                    values=list(self.IDIOMAS_WHISPER.keys()),
                                    state="readonly", width=20)
        idioma_combo.grid(row=row, column=1, padx=5, pady=5)

        # Mostrar nome do idioma selecionado
        self.idioma_label = ttk.Label(config_frame, text=self.IDIOMAS_WHISPER[self.idioma_escolhido.get()])
        self.idioma_label.grid(row=row, column=2, padx=5, pady=5)
        idioma_combo.bind("<<ComboboxSelected>>", self._atualizar_idioma_label)

        # Temperatura
        row += 1
        ttk.Label(config_frame, text="Temperatura (0.0-1.0):").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        temp_spin = ttk.Spinbox(config_frame, from_=0.0, to=1.0, increment=0.1,
                                textvariable=self.temperatura, width=10)
        temp_spin.grid(row=row, column=1, padx=5, pady=5, sticky="w")
        ttk.Label(config_frame, text="(menor = mais preciso, maior = mais criativo)",
                  foreground="gray").grid(row=row, column=2, padx=5, pady=5)

        # Duração do segmento
        row += 1
        ttk.Label(config_frame, text="Duração do Segmento (seg):").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        seg_spin = ttk.Spinbox(config_frame, from_=10, to=120, increment=10,
                               textvariable=self.segmento_duracao, width=10)
        seg_spin.grid(row=row, column=1, padx=5, pady=5, sticky="w")

        # Configurações de Saída
        saida_frame = ttk.LabelFrame(frame, text="Configurações de Saída", padding=10)
        saida_frame.pack(fill="x", padx=10, pady=10)

        # Pasta de saída personalizada
        ttk.Label(saida_frame, text="Pasta de Saída (opcional):").pack(anchor="w", pady=5)
        pasta_frame = ttk.Frame(saida_frame)
        pasta_frame.pack(fill="x", pady=5)

        pasta_entry = ttk.Entry(pasta_frame, textvariable=self.pasta_saida_personalizada, width=50)
        pasta_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        ttk.Button(pasta_frame, text="Procurar", command=self._selecionar_pasta_saida, width=10).pack(side="right")

        # Botões de configuração
        botoes_frame = ttk.Frame(frame)
        botoes_frame.pack(fill="x", padx=10, pady=20)

        ttk.Button(botoes_frame, text="Salvar Configurações", command=self._salvar_configuracoes).pack(side="left",
                                                                                                       padx=5)
        ttk.Button(botoes_frame, text="Restaurar Padrões", command=self._restaurar_padroes).pack(side="left", padx=5)

    def _configurar_aba_estatisticas(self):
        frame = self.frame_stats

        stats_frame = ttk.LabelFrame(frame, text="Estatísticas de Uso", padding=10)
        stats_frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.stats_text = Text(stats_frame, height=20, width=80, state=DISABLED, wrap="word")
        scrollbar_stats = ttk.Scrollbar(stats_frame, command=self.stats_text.yview)
        self.stats_text['yscrollcommand'] = scrollbar_stats.set

        self.stats_text.pack(side="left", fill="both", expand=True)
        scrollbar_stats.pack(side="right", fill="y")

        # Botão para atualizar estatísticas
        ttk.Button(frame, text="Atualizar Estatísticas", command=self._atualizar_estatisticas).pack(pady=10)

        self._atualizar_estatisticas()

    def _selecionar_pasta_saida(self):
        pasta = filedialog.askdirectory(title="Selecione a pasta de saída")
        if pasta:
            self.pasta_saida_personalizada.set(pasta)

    def _restaurar_padroes(self):
        self.modelo_escolhido.set(WhisperModel.TURBO.value)
        self.formato_saida.set("docx")
        self.idioma_escolhido.set("auto")
        self.temperatura.set("0.0")
        self.segmento_duracao.set(30)
        self.incluir_subpastas.set(False)
        self.incluir_timestamps.set(False)
        self.pasta_saida_personalizada.set("")
        self._atualizar_idioma_label()
        messagebox.showinfo("Sucesso", "Configurações restauradas para os valores padrão!")

    def _atualizar_idioma_label(self, event=None):
        idioma_selecionado = self.idioma_escolhido.get()
        nome_idioma = self.IDIOMAS_WHISPER.get(idioma_selecionado, "Desconhecido")
        self.idioma_label.config(text=nome_idioma)

    def _atualizar_descricao_modelo(self, event=None):
        descricao = self.MODELOS_DESCRICAO.get(self.modelo_escolhido.get(), "Descrição não disponível.")
        self.modelo_explicacao.config(text=descricao)
        # Limpar status do modelo carregado quando trocar
        if self.modelo_carregado_nome != self.modelo_escolhido.get():
            self.status_modelo.config(text="Modelo não carregado", foreground="red")

    def _atualizar_estatisticas(self):
        stats_text = f"""=== ESTATÍSTICAS DE USO ===

📊 Sessão Atual:
• Arquivos processados: {self.estatisticas['arquivos_processados']}
• Sucessos: {self.estatisticas['sucessos']}
• Erros: {self.estatisticas['erros']}
• Tempo total de processamento: {self._formatar_tempo(self.estatisticas['tempo_total_processamento'])}

⚙️ Configuração Atual:
• Modelo: {self.modelo_escolhido.get()}
• Idioma: {self.IDIOMAS_WHISPER.get(self.idioma_escolhido.get(), 'Desconhecido')}
• Formato de saída: {self.formato_saida.get().upper()}
• Temperatura: {self.temperatura.get()}
• Duração do segmento: {self.segmento_duracao.get()}s

🖥️ Sistema:
• OS: {platform.system()} {platform.release()}
• FFmpeg disponível: {'Sim' if PYDUB_AVAILABLE else 'Não'}
• Modelo carregado: {self.modelo_carregado_nome or 'Nenhum'}

📅 Última atualização: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
"""

        self.stats_text.config(state=NORMAL)
        self.stats_text.delete(1.0, "end")
        self.stats_text.insert("end", stats_text)
        self.stats_text.config(state=DISABLED)

    def _formatar_tempo(self, segundos):
        horas, resto = divmod(int(segundos), 3600)
        minutos, segs = divmod(resto, 60)
        return f"{horas:02d}:{minutos:02d}:{segs:02d}"

    def carregar_modelo(self):
        if not PYDUB_AVAILABLE:
            messagebox.showerror("Erro", "FFmpeg ou pydub não estão disponíveis. Não é possível carregar o modelo.")
            return None

        modelo_selecionado = self.modelo_escolhido.get()

        # Se o modelo já está carregado, retornar
        if self.modelo_carregado and self.modelo_carregado_nome == modelo_selecionado:
            return self.modelo_carregado

        logging.info(f"Carregando o modelo '{modelo_selecionado}'...")
        self._inserir_detalhes(f"🔄 Carregando modelo: {modelo_selecionado}...")
        self.status_modelo.config(text="Carregando modelo...", foreground="orange")
        self.root.update_idletasks()

        try:
            # Limpar modelo anterior
            self.modelo_carregado = None
            self.modelo_carregado_nome = None

            # Carregar novo modelo
            self.modelo_carregado = whisper.load_model(modelo_selecionado)
            self.modelo_carregado_nome = modelo_selecionado

            self.status_modelo.config(text=f"✅ Modelo {modelo_selecionado} carregado", foreground="green")
            self._inserir_detalhes(f"✅ Modelo carregado com sucesso: {modelo_selecionado}")

            return self.modelo_carregado

        except Exception as e:
            error_msg = f"Não foi possível carregar o modelo '{modelo_selecionado}'. Erro: {e}"
            messagebox.showerror("Erro de Carregamento", error_msg)
            logging.error(f"Erro ao carregar o modelo '{modelo_selecionado}': {e}")
            self._inserir_detalhes(f"❌ Falha ao carregar modelo: {modelo_selecionado} - {e}")
            self.status_modelo.config(text="❌ Erro ao carregar modelo", foreground="red")
            return None

    def iniciar_transcricao(self):
        if not PYDUB_AVAILABLE:
            messagebox.showerror("Erro", "FFmpeg ou pydub não estão disponíveis. Não é possível iniciar a transcrição.")
            return

        self._limpar_detalhes()
        tipos_arquivo = [("Arquivos de áudio", [ext.value for ext in AudioExtension]), ("Todos os arquivos", "*.*")]
        caminho_audio = filedialog.askopenfilename(title="Selecione o arquivo de áudio", filetypes=tipos_arquivo)

        if caminho_audio:
            modelo = self.carregar_modelo()
            if modelo is None:
                return

            self._set_transcription_controls_state(True)
            self.cancel_event.clear()
            self.pause_event.clear()
            self.total_bytes = os.path.getsize(caminho_audio)
            self.processed_bytes = 0
            self.start_time = time.time()
            threading.Thread(target=self.transcrever_audio, args=(modelo, caminho_audio, 1, 1), daemon=True).start()
        else:
            self.progresso_text_label.config(text="Nenhum arquivo selecionado.")

    def iniciar_transcricao_em_lote(self):
        if not PYDUB_AVAILABLE:
            messagebox.showerror("Erro",
                                 "FFmpeg ou pydub não estão disponíveis. Não é possível iniciar a transcrição em lote.")
            return

        self._limpar_detalhes()
        pasta = filedialog.askdirectory(title="Selecione a pasta com os arquivos de áudio")

        if pasta:
            arquivos_audio = self._selecionar_arquivos_audio(pasta)
            if not arquivos_audio:
                self.progresso_text_label.config(text="Nenhum arquivo de áudio encontrado na pasta selecionada.")
                return

            # Mostrar prévia dos arquivos encontrados
            preview = f"Encontrados {len(arquivos_audio)} arquivo(s):\n\n"
            for i, arquivo in enumerate(arquivos_audio[:10]):  # Mostrar apenas os primeiros 10
                preview += f"• {os.path.basename(arquivo)}\n"
            if len(arquivos_audio) > 10:
                preview += f"... e mais {len(arquivos_audio) - 10} arquivo(s)"

            if not messagebox.askyesno("Confirmar Transcrição em Lote", preview):
                return

            modelo = self.carregar_modelo()
            if modelo is None:
                return

            self._set_transcription_controls_state(True)
            self.cancel_event.clear()
            self.pause_event.clear()
            self.total_bytes = sum(os.path.getsize(arquivo) for arquivo in arquivos_audio)
            self.processed_bytes = 0
            self.start_time = time.time()
            threading.Thread(target=self.processar_em_lote, args=(modelo, arquivos_audio), daemon=True).start()
        else:
            self.progresso_text_label.config(text="Nenhuma pasta selecionada.")

    def _selecionar_arquivos_audio(self, pasta):
        arquivos_audio = []
        extensoes = [ext.value.lower() for ext in AudioExtension]

        if self.incluir_subpastas.get():
            for root_dir, _, files in os.walk(pasta):
                for f in files:
                    if any(f.lower().endswith(ext) for ext in extensoes):
                        arquivos_audio.append(os.path.join(root_dir, f))
        else:
            for f in os.listdir(pasta):
                if any(f.lower().endswith(ext) for ext in extensoes):
                    arquivos_audio.append(os.path.join(pasta, f))

        return sorted(arquivos_audio)  # Ordenar por nome

    def transcrever_audio(self, modelo, caminho_audio, indice, total, ultimo_arquivo=False):
        if modelo is None:
            self.progresso_text_label.config(text="Erro: Modelo Whisper não carregado.")
            self._set_transcription_controls_state(False)
            return

        arquivo_nome = os.path.basename(caminho_audio)
        pos_inicial = self._inserir_detalhes(f"🎵 Iniciando transcrição: {arquivo_nome}")
        segment_duration = self.segmento_duracao.get()
        transcricao_completa = ""
        temp_file_path = None
        arquivo_inicio = time.time()

        try:
            # Informações do arquivo
            tamanho_mb = os.path.getsize(caminho_audio) / (1024 * 1024)
            self._inserir_detalhes(f"📄 Arquivo: {arquivo_nome} ({tamanho_mb:.1f} MB)")

            audio = AudioSegment.from_file(caminho_audio)
            duration = len(audio) / 1000  # Duração total em segundos
            segments_count = int(duration / segment_duration) + (1 if duration % segment_duration > 0 else 0)

            self._inserir_detalhes(f"⏱️ Duração: {self._formatar_tempo(duration)} | Segmentos: {segments_count}")

            self.progresso_barra['value'] = 0
            self.progresso_text_label.config(text=f"Transcrevendo: {arquivo_nome}...")

            # Configurações do Whisper
            idioma = None if self.idioma_escolhido.get() == "auto" else self.idioma_escolhido.get()
            temperatura = float(self.temperatura.get())

            for i, start_time_sec in enumerate(range(0, int(duration), segment_duration)):
                if self.cancel_event.is_set():
                    self.progresso_text_label.config(text="Transcrição cancelada.")
                    self._substituir_detalhes(pos_inicial, f"❌ Transcrição cancelada: {arquivo_nome}")
                    return

                while self.pause_event.is_set():
                    self.root.update_idletasks()
                    time.sleep(0.1)

                end_time_sec = min(start_time_sec + segment_duration, duration)
                segment = audio[start_time_sec * 1000:end_time_sec * 1000]

                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
                    temp_file_path = temp_file.name
                    segment.export(temp_file_path, format="wav")

                # Transcreve o segmento com configurações avançadas
                result = modelo.transcribe(
                    temp_file_path,
                    language=idioma,
                    temperature=temperatura,
                    task="transcribe"
                )

                # Adicionar timestamps se solicitado
                if self.incluir_timestamps.get():
                    timestamp = f"[{self._formatar_tempo(start_time_sec)} -> {self._formatar_tempo(end_time_sec)}] "
                    transcricao_completa += timestamp + result["text"] + "\n\n"
                else:
                    transcricao_completa += result["text"] + " "

                # Remove o arquivo temporário imediatamente
                if temp_file_path and os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                    temp_file_path = None

                # Atualizar progresso
                progresso_segmento = ((i + 1) / segments_count) * 100
                self.progresso_barra['value'] = progresso_segmento

                # Calcular ETA
                elapsed = time.time() - arquivo_inicio
                if i > 0:
                    eta_total = (elapsed / (i + 1)) * segments_count - elapsed
                    self.eta_label.config(text=f"TEMPO RESTANTE: {self._formatar_tempo(eta_total)}")

                self.root.update_idletasks()

        except FileNotFoundError:
            messagebox.showerror("Erro", f"Arquivo não encontrado: {arquivo_nome}")
            logging.error(f"Arquivo não encontrado: {caminho_audio}")
            self._substituir_detalhes(pos_inicial, f"❌ Erro: Arquivo não encontrado: {arquivo_nome}")
            self.estatisticas['erros'] += 1
        except Exception as e:
            messagebox.showerror("Erro de Transcrição", f"Erro ao transcrever '{arquivo_nome}'. Erro: {e}")
            logging.error(f"Erro na transcrição de '{caminho_audio}': {e}", exc_info=True)
            self._substituir_detalhes(pos_inicial, f"❌ Erro na transcrição: {arquivo_nome} - {e}")
            self.estatisticas['erros'] += 1
        finally:
            # Garantir limpeza do arquivo temporário
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except:
                    pass

            if not self.cancel_event.is_set():
                if transcricao_completa.strip():
                    tempo_arquivo = time.time() - arquivo_inicio
                    self.salvar_transcricao(transcricao_completa, caminho_audio, self.formato_saida.get())
                    self.processed_bytes += os.path.getsize(caminho_audio)
                    self._atualizar_progresso(indice, total)
                    self._substituir_detalhes(pos_inicial,
                                              f"✅ Transcrito com sucesso: {arquivo_nome} ({self._formatar_tempo(tempo_arquivo)})")
                    self.estatisticas['sucessos'] += 1
                    self.estatisticas['tempo_total_processamento'] += tempo_arquivo
                else:
                    self.progresso_text_label.config(text=f"Transcrição vazia para {arquivo_nome}. Verifique o áudio.")
                    self._substituir_detalhes(pos_inicial, f"⚠️ Transcrição vazia: {arquivo_nome}")

            self.estatisticas['arquivos_processados'] += 1

            # Verificar se deve reabilitar controles
            if ultimo_arquivo or self.cancel_event.is_set() or indice == total:
                self.progresso_text_label.config(text="Transcrição concluída! Pronto para nova transcrição.")
                self.eta_label.config(text="")
                self._set_transcription_controls_state(False)

    def _atualizar_progresso(self, indice, total):
        elapsed_time = time.time() - self.start_time
        progresso_percentual = (self.processed_bytes / self.total_bytes) * 100 if self.total_bytes > 0 else 0
        self.progresso_barra['value'] = progresso_percentual

        # ETA para processamento em lote
        if indice > 0 and total > 1:
            eta_restante = (elapsed_time / indice) * (total - indice)
            eta_text = f"ETA total: {self._formatar_tempo(eta_restante)}"
        else:
            eta_text = ""

        progresso_text = f"{progresso_percentual:.0f}% [{indice}/{total}] Tempo: {self._formatar_tempo(elapsed_time)}"
        self.progresso_text_label.config(text=progresso_text)
        self.eta_label.config(text=eta_text)
        self.root.update_idletasks()

    def salvar_transcricao(self, texto_transcrito, caminho_audio, formato='docx'):
        # Determinar pasta de saída
        if self.pasta_saida_personalizada.get():
            pasta_saida = self.pasta_saida_personalizada.get()
        else:
            pasta_saida = os.path.dirname(caminho_audio)

        nome_arquivo = os.path.splitext(os.path.basename(caminho_audio))[0]
        modelo = self.modelo_escolhido.get()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if formato == 'srt':
            # Para SRT, precisamos dos timestamps
            caminho_saida = os.path.join(pasta_saida, f"{nome_arquivo}_transcrito_{modelo}_{timestamp}.srt")
        else:
            caminho_saida = os.path.join(pasta_saida, f"{nome_arquivo}_transcrito_{modelo}_{timestamp}.{formato}")

        try:
            # Criar pasta se não existir
            os.makedirs(pasta_saida, exist_ok=True)

            if formato == 'txt':
                with open(caminho_saida, 'w', encoding='utf-8') as f:
                    header = f"=== TRANSCRIÇÃO DE ÁUDIO ===\n"
                    header += f"Arquivo: {os.path.basename(caminho_audio)}\n"
                    header += f"Modelo: {modelo}\n"
                    header += f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n"
                    header += f"Idioma: {self.IDIOMAS_WHISPER.get(self.idioma_escolhido.get(), 'Auto')}\n"
                    header += "=" * 50 + "\n\n"
                    f.write(header + texto_transcrito)

            elif formato == 'markdown':
                with open(caminho_saida, 'w', encoding='utf-8') as f:
                    header = f"# Transcrição de {nome_arquivo}\n\n"
                    header += f"**Modelo:** {modelo}  \n"
                    header += f"**Data:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}  \n"
                    header += f"**Idioma:** {self.IDIOMAS_WHISPER.get(self.idioma_escolhido.get(), 'Auto')}  \n\n"
                    header += "---\n\n## Conteúdo\n\n"
                    f.write(header + texto_transcrito)

            elif formato == 'srt':
                with open(caminho_saida, 'w', encoding='utf-8') as f:
                    # Converter para formato SRT básico
                    if self.incluir_timestamps.get():
                        f.write(texto_transcrito)
                    else:
                        f.write("1\n00:00:00,000 --> 99:59:59,999\n" + texto_transcrito.strip() + "\n")

            else:  # docx
                doc = Document()
                doc.add_heading(f"Transcrição de {nome_arquivo}", level=1)

                # Adicionar metadados
                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Table Grid'

                cells = info_table.rows[0].cells
                cells[0].text = "Modelo Whisper"
                cells[1].text = modelo

                cells = info_table.rows[1].cells
                cells[0].text = "Data da Transcrição"
                cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

                cells = info_table.rows[2].cells
                cells[0].text = "Idioma"
                cells[1].text = self.IDIOMAS_WHISPER.get(self.idioma_escolhido.get(), 'Auto')

                cells = info_table.rows[3].cells
                cells[0].text = "Arquivo Original"
                cells[1].text = os.path.basename(caminho_audio)

                doc.add_paragraph("")  # Espaço
                doc.add_heading("Conteúdo da Transcrição", level=2)
                doc.add_paragraph(texto_transcrito)
                doc.save(caminho_saida)

            logging.info(f"Transcrição salva no arquivo: {caminho_saida}")

            # Perguntar se quer abrir a pasta apenas no final da transcrição individual
            # ou no último arquivo do lote
            if messagebox.askyesno("Transcrição Concluída",
                                   f"Transcrição de '{nome_arquivo}' salva com sucesso!\n\n"
                                   f"Local: {caminho_saida}\n\n"
                                   f"Deseja abrir a pasta onde o arquivo foi salvo?"):
                self.abrir_pasta(pasta_saida)

        except Exception as e:
            messagebox.showerror("Erro ao Salvar",
                                 f"Não foi possível salvar a transcrição para '{nome_arquivo}'. Erro: {e}")
            logging.error(f"Erro ao salvar transcrição para '{caminho_audio}': {e}")

    def abrir_pasta(self, caminho_pasta):
        try:
            if platform.system() == "Windows":
                os.startfile(caminho_pasta)
            elif platform.system() == "Darwin":  # macOS
                subprocess.Popen(["open", caminho_pasta])
            else:  # Linux e outros sistemas Unix-like
                subprocess.Popen(["xdg-open", caminho_pasta])
        except Exception as e:
            logging.error(f"Falha ao abrir a pasta '{caminho_pasta}': {e}")
            messagebox.showerror("Erro", f"Não foi possível abrir a pasta: {e}")

    def processar_em_lote(self, modelo, arquivos_audio):
        if modelo is None:
            self.progresso_text_label.config(text="Erro: Modelo Whisper não carregado para transcrição em lote.")
            self._set_transcription_controls_state(False)
            return

        total = len(arquivos_audio)
        inicio_lote = time.time()
        self._inserir_detalhes(f"🚀 Iniciando processamento em lote de {total} arquivo(s)")

        for index, caminho_audio in enumerate(arquivos_audio):
            if self.cancel_event.is_set():
                self.progresso_text_label.config(text="Processo de lote cancelado.")
                break

            while self.pause_event.is_set():
                self.root.update_idletasks()
                time.sleep(0.1)

            ultimo_arquivo = index == total - 1
            self.transcrever_audio(modelo, caminho_audio, index + 1, total, ultimo_arquivo=ultimo_arquivo)

        # Resumo final
        tempo_total = time.time() - inicio_lote
        if not self.cancel_event.is_set():
            self.progresso_text_label.config(
                text=f"Lote concluído! {self.estatisticas['sucessos']} sucessos, {self.estatisticas['erros']} erros.")
            self._inserir_detalhes(f"🎉 Processamento em lote concluído em {self._formatar_tempo(tempo_total)}")
            self._inserir_detalhes(
                f"📊 Resumo: {self.estatisticas['sucessos']} sucessos, {self.estatisticas['erros']} erros")

        self._set_transcription_controls_state(False)

    def cancelar_processo(self):
        self.cancel_event.set()
        logging.info("Sinal de cancelamento enviado para o processo de transcrição.")
        self.progresso_text_label.config(text="Cancelando transcrição...")
        self.botao_pausar.config(text="Pausar")
        self.pause_event.clear()

    def pausar_processo(self):
        if self.pause_event.is_set():
            self.pause_event.clear()
            novo_texto = "Pausar"
            logging.info("Processo de transcrição retomado.")
            self.progresso_text_label.config(text="Transcrição retomada.")
        else:
            self.pause_event.set()
            novo_texto = "Continuar"
            logging.info("Processo de transcrição pausado.")
            self.progresso_text_label.config(text="Transcrição pausada.")

        self.botao_pausar.config(text=novo_texto)

    def _set_transcription_controls_state(self, transcricao_ativa: bool):
        """Controla o estado dos botões durante a transcrição"""
        if transcricao_ativa:
            # Durante transcrição: desabilita iniciar, habilita cancelar/pausar
            self.btn_transcricao_individual.config(state=DISABLED)
            self.btn_transcricao_em_lote.config(state=DISABLED)
            self.botao_cancelar.config(state=NORMAL)
            self.botao_pausar.config(state=NORMAL)
        else:
            # Transcrição finalizada: habilita iniciar, desabilita cancelar/pausar
            self.btn_transcricao_individual.config(state=NORMAL)
            self.btn_transcricao_em_lote.config(state=NORMAL)
            self.botao_cancelar.config(state=DISABLED)
            self.botao_pausar.config(state=DISABLED)
            # Reset do botão pausar
            self.botao_pausar.config(text="Pausar")
            self.pause_event.clear()

    def _inserir_detalhes(self, mensagem):
        self.detalhes_text.config(state=NORMAL)
        pos_inicial = self.detalhes_text.index("end-1c")
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.detalhes_text.insert("end", f"[{timestamp}] {mensagem}\n")
        self.detalhes_text.config(state=DISABLED)
        self.detalhes_text.see("end")
        return pos_inicial

    def _substituir_detalhes(self, pos_inicial, mensagem):
        self.detalhes_text.config(state=NORMAL)
        try:
            self.detalhes_text.delete(pos_inicial, f"{pos_inicial} lineend")
        except:
            pass
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.detalhes_text.insert(pos_inicial, f"[{timestamp}] {mensagem}\n")
        self.detalhes_text.config(state=DISABLED)
        self.detalhes_text.see("end")

    def _alternar_painel_detalhes(self):
        if self.painel_detalhes.winfo_ismapped():
            self.painel_detalhes.grid_remove()
            self.botao_detalhes.config(text="Ver Detalhes")
        else:
            self.painel_detalhes.grid(row=18, column=0, columnspan=2, padx=10, pady=5, sticky="nsew")
            self.botao_detalhes.config(text="Ocultar Detalhes")

    def _limpar_detalhes(self):
        self.detalhes_text.config(state=NORMAL)
        self.detalhes_text.delete(1.0, "end")
        self.detalhes_text.config(state=DISABLED)

    def _fechar_aplicacao(self):
        # Salvar configurações antes de fechar
        self._salvar_configuracoes()

        # Se houver transcrição em andamento, perguntar se quer cancelar
        if str(self.botao_cancelar.cget('state')) != DISABLED:
            if messagebox.askyesno("Fechar Aplicação",
                                   "Há uma transcrição em andamento. Deseja cancelar e fechar?"):
                self.cancel_event.set()
            else:
                return

        self.root.quit()

    def iniciar_interface(self):
        # Configurar evento de fechamento da janela
        self.root.protocol("WM_DELETE_WINDOW", self._fechar_aplicacao)
        self.root.mainloop()


if __name__ == "__main__":
    app = TranscricaoAudio()
    app.iniciar_interface()